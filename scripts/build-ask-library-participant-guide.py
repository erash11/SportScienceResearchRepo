from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = (
    ROOT
    / "docs"
    / "ask-library-pilot"
    / "Ask-the-Library-Pilot-Participant-Quick-Start.docx"
)

GREEN = "154734"
DARK_GREEN = "0B2B21"
GOLD = "FFB81C"
INK = "202722"
MUTED = "5E6B64"
SOFT_GREEN = "E8F0EC"
SOFT_GOLD = "FFF4D6"
WHITE = "FFFFFF"


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, *, side="left", color=GOLD, size=20, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def add_inline_paragraph(doc, label, text, *, after=4):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=10.5, color=DARK_GREEN, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return paragraph


def create_decimal_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_alignment = OxmlElement("w:lvlJc")
    level_alignment.set(qn("w:val"), "left")
    level.append(level_alignment)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), "Calibri")
    run_fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(run_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GOLD)
    run_properties.append(color)
    bold = OxmlElement("w:b")
    run_properties.append(bold)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return abstract_id, num_id


def create_number_instance(document, abstract_id):
    numbering = document.part.numbering_part.element
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(num_ids, default=0) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def add_numbered_item(doc, num_id, label, text):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(number)
    p_pr.insert(0, num_pr)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=10.5, color=DARK_GREEN, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    return paragraph


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor.from_string(GREEN)
    heading_1.paragraph_format.space_before = Pt(12)
    heading_1.paragraph_format.space_after = Pt(6)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True


def build_document():
    document = Document()
    document.core_properties.title = "Ask the Library Pilot Participant Quick Start"
    document.core_properties.subject = "Participant instructions for the concierge pilot"
    document.core_properties.author = "Baylor Athletics"
    document.core_properties.keywords = "Ask the Library, pilot, Decision Brief"

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.85)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.start_type = WD_SECTION.NEW_PAGE
    configure_styles(document)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(
        "BAYLOR ATHLETICS  |  HEALTH & PERFORMANCE EVIDENCE"
    )
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(4)
    kicker.paragraph_format.space_after = Pt(3)
    kicker_run = kicker.add_run("CONCIERGE PILOT  /  PARTICIPANT QUICK START")
    set_run_font(kicker_run, size=9, color=GOLD, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("Ask the Library")
    set_run_font(title_run, size=27, color=DARK_GREEN, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.1
    subtitle_run = subtitle.add_run(
        "Bring one real decision. Leave with a short, source-grounded brief."
    )
    set_run_font(subtitle_run, size=12.5, color=GREEN, bold=True)

    intro = document.add_paragraph()
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(10)
    intro.paragraph_format.left_indent = Inches(0.16)
    intro.paragraph_format.right_indent = Inches(0.16)
    intro.paragraph_format.line_spacing = 1.1
    set_paragraph_shading(intro, GREEN)
    set_paragraph_border(intro, side="left", color=GOLD, size=24, space=6)
    intro_run = intro.add_run(
        "Your role is simple: ask, read, and rate. Enter one de-identified "
        "question, read the finished answer, and complete about one minute of "
        "feedback. The pilot lead handles Codex and all files."
    )
    set_run_font(intro_run, size=10.5, color=WHITE, bold=True)

    document.add_heading("Before each session", level=1)
    before_abstract, before_num = create_decimal_numbering(document)
    add_numbered_item(
        document,
        before_num,
        "Choose a real decision.",
        "Bring something you are planning, discussing, monitoring, or deciding now.",
    )
    add_numbered_item(
        document,
        before_num,
        "Remove identifying details.",
        "Do not include names, initials, roster numbers, medical records, clinical notes, dates of birth, or contact information.",
    )
    add_numbered_item(
        document,
        before_num,
        "Bring the context that matters.",
        "Be ready to describe the population, sport or setting, phase or timing, intended outcome, and operational constraints.",
    )

    question = document.add_paragraph()
    question.paragraph_format.space_before = Pt(7)
    question.paragraph_format.space_after = Pt(4)
    question.paragraph_format.left_indent = Inches(0.16)
    question.paragraph_format.right_indent = Inches(0.16)
    question.paragraph_format.line_spacing = 1.1
    set_paragraph_shading(question, SOFT_GOLD)
    set_paragraph_border(question, side="left", color=GOLD, size=20, space=6)
    question_label = question.add_run("QUESTION FORMAT  ")
    set_run_font(question_label, size=9.5, color=GREEN, bold=True)
    question_text = question.add_run(
        "For [population] in [phase or setting], how should we approach "
        "[decision] to support [outcome], given [constraints]?"
    )
    set_run_font(question_text, size=10, color=INK, italic=True)

    example = document.add_paragraph()
    example.paragraph_format.space_before = Pt(0)
    example.paragraph_format.space_after = Pt(5)
    example.paragraph_format.left_indent = Inches(0.16)
    example.paragraph_format.right_indent = Inches(0.16)
    example.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(example, SOFT_GOLD)
    example_label = example.add_run("Example: ")
    set_run_font(example_label, size=9.5, color=DARK_GREEN, bold=True)
    example_text = example.add_run(
        "How should we adjust field exposure for collegiate starters between "
        "competitions inside 96 hours when travel limits recovery time?"
    )
    set_run_font(example_text, size=9.5, color=INK)

    document.add_heading("During the session: ask, read, rate", level=1)
    during_num = create_number_instance(document, before_abstract)
    add_numbered_item(
        document,
        during_num,
        "Ask.",
        "Copy an assigned prompt or enter another de-identified question, add the context that matters, and select Save question for Codex. The pilot lead handles the downloaded file.",
    )
    add_numbered_item(
        document,
        during_num,
        "Read.",
        "The pilot lead asks Codex to prepare and check the answer, then opens the finished brief. Start with Bottom Line, Recommended Direction or Decision Boundary, Evidence Confidence, and Guardrails.",
    )
    add_numbered_item(
        document,
        during_num,
        "Rate.",
        "Complete the short feedback form: Did it help? Did it affect the decision? What was missing? How long did it take to understand?",
    )

    role_note = document.add_paragraph()
    role_note.paragraph_format.space_before = Pt(5)
    role_note.paragraph_format.space_after = Pt(4)
    role_note.paragraph_format.left_indent = Inches(0.16)
    role_note.paragraph_format.right_indent = Inches(0.16)
    role_note.paragraph_format.line_spacing = 1.1
    set_paragraph_shading(role_note, SOFT_GREEN)
    role_label = role_note.add_run("THAT IS YOUR WHOLE ROLE  ")
    set_run_font(role_label, size=9, color=GREEN, bold=True)
    role_text = role_note.add_run(
        "You do not need to use Codex, manage JSON files, open GitHub, or run commands."
    )
    set_run_font(role_text, size=9.5, color=INK, bold=True)

    document.add_heading("How to interpret the result", level=1)
    add_inline_paragraph(
        document,
        "On-Demand / Not Expert-Reviewed.",
        "The brief was prepared for this pilot and has not been promoted to a formally reviewed Baylor resource.",
        after=2,
    )
    add_inline_paragraph(
        document,
        "Coverage Gap is valid.",
        "It means the library does not currently support a trustworthy direction.",
        after=2,
    )
    add_inline_paragraph(
        document,
        "Staff judgment stays primary.",
        "The brief is not a diagnosis, clearance decision, prescription, or Baylor policy.",
        after=4,
    )

    privacy = document.add_paragraph()
    privacy.paragraph_format.space_before = Pt(4)
    privacy.paragraph_format.space_after = Pt(0)
    privacy.paragraph_format.line_spacing = 1.1
    privacy.paragraph_format.left_indent = Inches(0.16)
    privacy.paragraph_format.right_indent = Inches(0.16)
    set_paragraph_shading(privacy, SOFT_GREEN)
    set_paragraph_border(privacy, side="left", color=GREEN, size=18, space=6)
    privacy_label = privacy.add_run("PRIVACY STOP RULE  ")
    set_run_font(privacy_label, size=9, color=GREEN, bold=True)
    privacy_text = privacy.add_run(
        "If identifying information appears, stop and resubmit the question using group-level or otherwise de-identified context."
    )
    set_run_font(privacy_text, size=9.5, color=INK, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run(
        "ASK THE LIBRARY  |  14-DAY MINIMUM CREDIBLE PILOT"
    )
    set_run_font(footer_run, size=8, color=MUTED, bold=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
