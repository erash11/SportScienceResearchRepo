import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "docs"
    / "ask-library-pilot"
    / "examples"
    / "stress-test-question-bank.json"
)
OUTPUT = (
    ROOT
    / "docs"
    / "ask-library-pilot"
    / "Ask-the-Library-Pilot-Stress-Test-Question-Bank.docx"
)

# compact_reference_guide with named Baylor green-and-gold override.
GREEN = "154734"
DARK_GREEN = "0B2B21"
GOLD = "FFB81C"
INK = "202722"
MUTED = "5E6B64"
SOFT_GREEN = "E8F0EC"
SOFT_GOLD = "FFF4D6"
LIGHT_GRAY = "F2F4F3"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, *, side="left", color=GOLD, size=20, space=8):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    existing = borders.find(qn(f"w:{side}"))
    if existing is not None:
        borders.remove(existing)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    borders.append(border)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA}: {widths}")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color="B9C5BF", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor.from_string(GREEN)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor.from_string(GREEN)
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(5)
    heading_3.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=8, color=MUTED, bold=True)


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION.NEW_PAGE

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_run = header_paragraph.add_run(
        "BAYLOR ATHLETICS  |  HEALTH & PERFORMANCE EVIDENCE"
    )
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    label = footer_paragraph.add_run("ASK THE LIBRARY PILOT  |  ")
    set_run_font(label, size=8, color=MUTED, bold=True)
    add_page_field(footer_paragraph)


def add_kicker(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=9, color=GOLD, bold=True)
    return paragraph


def add_numbered_step(document, number, label, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    number_run = paragraph.add_run(f"{number}. ")
    set_run_font(number_run, size=10.5, color=GOLD, bold=True)
    label_run = paragraph.add_run(f"{label} ")
    set_run_font(label_run, size=10.5, color=DARK_GREEN, bold=True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, size=10.5, color=INK)


def add_assignment_table(document):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1800, 3540, 4020])
    set_table_borders(table)

    headers = ["Domain track", "Best-fit practitioner", "Questions"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, GREEN)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=9.5, color=WHITE, bold=True)
    mark_header_row(table.rows[0])

    rows = [
        (
            "Physical Therapy",
            "Physical therapist / rehabilitation lead",
            "PT01 Simple · PT02 Applied · PT03 Complex",
        ),
        (
            "Performance Nutrition",
            "Sports dietitian / performance nutritionist",
            "PN01 Simple · PN02 Applied · PN03 Complex",
        ),
        (
            "Sports Science",
            "Sports scientist / performance analyst",
            "SS01 Simple · SS02 Applied · SS03 Complex",
        ),
        (
            "Sports Performance",
            "Strength and conditioning / performance coach",
            "SP01 Simple · SP02 Applied · SP03 Complex",
        ),
    ]
    for row_number, values in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(values):
            set_cell_margins(cells[index])
            if row_number % 2 == 1:
                set_cell_shading(cells[index], LIGHT_GRAY)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(text)
            set_run_font(
                run,
                size=9.25,
                color=DARK_GREEN if index == 0 else INK,
                bold=index == 0,
            )
    set_table_geometry(table, [1800, 3540, 4020])
    return table


def add_callout(document, label, text, *, fill=SOFT_GOLD, border=GOLD):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.14)
    paragraph.paragraph_format.line_spacing = 1.12
    set_paragraph_shading(paragraph, fill)
    set_paragraph_border(paragraph, side="left", color=border, size=20, space=6)
    label_run = paragraph.add_run(f"{label.upper()}  ")
    set_run_font(label_run, size=9, color=DARK_GREEN, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=9.75, color=INK, bold=True)
    return paragraph


def add_context_line(document, values):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, (label, text) in enumerate(values):
        if index:
            separator = paragraph.add_run("  ·  ")
            set_run_font(separator, size=8.8, color=MUTED)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=8.8, color=DARK_GREEN, bold=True)
        value_run = paragraph.add_run(text)
        set_run_font(value_run, size=8.8, color=INK)
    return paragraph


def add_question_card(document, question):
    tier = question["tier"].upper()
    title = f'{question["questionId"]} · {tier}  |  {question["shortTitle"]}'
    document.add_heading(title, level=3)

    prompt = document.add_paragraph()
    prompt.paragraph_format.space_before = Pt(0)
    prompt.paragraph_format.space_after = Pt(4)
    prompt.paragraph_format.left_indent = Inches(0.13)
    prompt.paragraph_format.right_indent = Inches(0.13)
    prompt.paragraph_format.line_spacing = 1.08
    set_paragraph_shading(prompt, SOFT_GREEN)
    set_paragraph_border(
        prompt,
        side="left",
        color=GOLD if question["tier"] == "complex" else GREEN,
        size=18,
        space=6,
    )
    run = prompt.add_run(question["practicalQuestion"])
    set_run_font(run, size=10.25, color=DARK_GREEN, bold=True)

    context = question["decisionContext"]
    add_context_line(
        document,
        [
            ("Population", context["population"]),
            ("Sport / setting", context["sport"]),
            ("Phase", context["phase"]),
        ],
    )
    add_context_line(
        document,
        [
            ("Outcome", context["outcome"]),
            ("Constraints", context["constraints"]),
        ],
    )

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(0)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.0
    label_run = note.add_run("PILOT LEAD ONLY — INTENDED STRESS  ")
    set_run_font(label_run, size=8, color=GOLD, bold=True)
    text_run = note.add_run(question["pilotLeadOnly"]["intendedStress"])
    set_run_font(text_run, size=8.25, color=MUTED, italic=True)


def add_domain_page(document, discipline, questions):
    document.add_page_break()
    add_kicker(document, f"{discipline} track / simple → applied → complex")

    # Named domain-page-title override: Heading 2 semantics with larger display type.
    title = document.add_paragraph(style="Heading 2")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(discipline)
    set_run_font(run, size=23, color=DARK_GREEN, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(5)
    subtitle.paragraph_format.line_spacing = 1.1
    run = subtitle.add_run(
        "Copy the Practical Question and all five Decision Context fields into the prototype."
    )
    set_run_font(run, size=10.5, color=GREEN, bold=True)

    check = document.add_paragraph()
    check.paragraph_format.space_before = Pt(0)
    check.paragraph_format.space_after = Pt(4)
    check.paragraph_format.line_spacing = 1.0
    label = check.add_run("RELEVANCE CHECK  ")
    set_run_font(label, size=8.5, color=GOLD, bold=True)
    body = check.add_run(
        "Participant confirms each scenario maps to a current or recent decision. If not, replace it at the same tier."
    )
    set_run_font(body, size=8.75, color=MUTED, italic=True)

    for question in questions:
        add_question_card(document, question)


def audit_geometry(document):
    section = document.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    assert round(section.top_margin.inches, 3) == 1.0
    assert round(section.right_margin.inches, 3) == 1.0
    assert round(section.bottom_margin.inches, 3) == 1.0
    assert round(section.left_margin.inches, 3) == 1.0
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492

    normal = document.styles["Normal"]
    assert normal.font.name == "Calibri"
    assert round(normal.font.size.pt, 1) == 11.0
    assert normal.paragraph_format.line_spacing == 1.25

    heading_values = {
        "Heading 1": (16, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (12, 10, 5),
    }
    for name, (size, before, after) in heading_values.items():
        style = document.styles[name]
        assert round(style.font.size.pt, 1) == size
        assert round(style.paragraph_format.space_before.pt, 1) == before
        assert round(style.paragraph_format.space_after.pt, 1) == after

    for table in document.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == str(TABLE_WIDTH_DXA)
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == str(
            TABLE_INDENT_DXA
        )


def build_document():
    bank = json.loads(SOURCE.read_text(encoding="utf-8"))
    questions = bank["questions"]
    if len(questions) != 12:
        raise ValueError("The four domain tracks require exactly twelve questions.")

    expected_tiers = {"simple", "applied", "complex"}
    expected_disciplines = {
        "Physical Therapy",
        "Performance Nutrition",
        "Sports Science",
        "Sports Performance",
    }
    grouped = {}
    for question in questions:
        grouped.setdefault(question["discipline"], []).append(question)
    if set(grouped) != expected_disciplines:
        raise ValueError(f"Unexpected domain tracks: {set(grouped)}")
    for discipline, items in grouped.items():
        tiers = {item["tier"] for item in items}
        if tiers != expected_tiers:
            raise ValueError(f"{discipline} does not contain all three tiers: {tiers}")

    document = Document()
    document.core_properties.title = bank["title"]
    document.core_properties.subject = (
        "Assignment-ready questions for the Ask the Library minimum credible pilot"
    )
    document.core_properties.author = "Baylor Athletics"
    document.core_properties.keywords = (
        "Ask the Library, pilot, stress test, Decision Brief, practitioner evaluation"
    )

    configure_styles(document)
    configure_section(document.sections[0])

    add_kicker(document, "Concierge pilot / question assignment pack")

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Ask the Library")
    set_run_font(run, size=29, color=DARK_GREEN, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.05
    run = subtitle.add_run("Pilot Stress-Test Question Bank")
    set_run_font(run, size=15, color=GREEN, bold=True)

    intro = document.add_paragraph()
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(10)
    intro.paragraph_format.left_indent = Inches(0.15)
    intro.paragraph_format.right_indent = Inches(0.15)
    intro.paragraph_format.line_spacing = 1.12
    set_paragraph_shading(intro, GREEN)
    set_paragraph_border(intro, side="left", color=GOLD, size=24, space=6)
    run = intro.add_run(
        "Twelve ready-to-enter scenarios cover four practitioner domains. For the current "
        "three-person pilot, select three matching tracks—nine submitted questions total. "
        "Each participant receives one simple, one applied, and one complex decision."
    )
    set_run_font(run, size=10.5, color=WHITE, bold=True)

    document.add_heading("How to run the stress test", level=1)
    add_numbered_step(
        document,
        1,
        "Select three domain tracks.",
        "Match one track to each participant; keep the fourth as an alternate so the pilot remains nine questions.",
    )
    add_numbered_step(
        document,
        2,
        "Copy and confirm the question.",
        "Enter the question and all five context fields, confirm real-world relevance, then select Save question for Codex.",
    )
    add_numbered_step(
        document,
        3,
        "Evaluate the brief, not the prompt.",
        "Use the existing feedback ledger and ask what is useful, wrong, missing, misapplied, or potentially unsafe.",
    )

    document.add_heading("Assignment matrix", level=1)
    add_assignment_table(document)

    add_callout(
        document,
        "Practitioner evaluation prompt",
        bank["evaluationPrompt"],
        fill=SOFT_GREEN,
        border=GREEN,
    )
    add_callout(
        document,
        "File rule",
        bank["usageNote"],
        fill=SOFT_GREEN,
        border=GREEN,
    )

    reminder = document.add_paragraph()
    reminder.paragraph_format.space_before = Pt(2)
    reminder.paragraph_format.space_after = Pt(0)
    reminder.paragraph_format.line_spacing = 1.05
    label = reminder.add_run("PILOT LEAD NOTE  ")
    set_run_font(label, size=8.5, color=GOLD, bold=True)
    body = reminder.add_run(
        "The intended-stress notes are not answer keys. Do not use them to steer the synthesis operator toward a preferred conclusion."
    )
    set_run_font(body, size=8.75, color=MUTED, italic=True)

    domain_assignments = [
        "Physical Therapy",
        "Performance Nutrition",
        "Sports Science",
        "Sports Performance",
    ]
    for discipline in domain_assignments:
        add_domain_page(document, discipline, grouped[discipline])

    audit_geometry(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
